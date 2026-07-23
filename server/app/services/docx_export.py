"""Render format Blocks into a Word document.

Kept separate from services/formats.py so the emitters stay free of docx
imports and can be tested as plain strings.

Layout mirrors the requested outline:

    # CDRs
    ## Telenor CDR
        <sets>
    ## Mobilink CDR
        ...
    # IMEIs
    ## Telenor IMEI
        ...

Networks appear in a fixed order (Telenor, Mobilink, Ufone, Zong); a network with
no rows in a section is skipped rather than shown empty. Gateway and any
unclassifiable rows follow in their own top-level sections.
"""

from __future__ import annotations

from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from .formats import SECTION_NETWORKS, Block, build_document

_MONO = "Consolas"

#: One distinct colour per network subheading, so sections are told apart fast.
_NETWORK_COLORS = {
    "telenor": RGBColor(0x1F, 0x4E, 0x79),   # blue
    "ufone": RGBColor(0x1B, 0x7F, 0x3B),     # green
    "mobilink": RGBColor(0xB0, 0x2A, 0x2A),  # red
    "zong": RGBColor(0x6A, 0x1B, 0x9A),      # purple
}
_DEFAULT_HEADING_COLOR = RGBColor(0x8A, 0x4B, 0x00)  # amber - gateway / unknown
_SECTION_COLOR = RGBColor(0x11, 0x11, 0x11)          # near-black - the # headings
_SUBTITLE_COLOR = RGBColor(0x0B, 0x61, 0x6E)         # teal - the date-range line
_NOTE_COLOR = RGBColor(0xB0, 0x2A, 0x2A)             # red - block-level notes

_NETWORK_NAMES = {
    "telenor": "Telenor",
    "mobilink": "Mobilink",
    "ufone": "Ufone",
    "zong": "Zong",
}


def _para(doc, *, before: float = 0, after: float = 0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


def _heading(doc, text: str, *, size: int, color: RGBColor, before: float, after: float):
    run = _para(doc, before=before, after=after).add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color


def _subtitle(doc, block: Block) -> None:
    if not (block.subtitle or block.note):
        return
    p = _para(doc, after=6)
    if block.subtitle:
        r = p.add_run(block.subtitle)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = _SUBTITLE_COLOR
    if block.note:
        if block.subtitle:
            p.add_run("    ").font.size = Pt(9)
        r = p.add_run(block.note)
        r.font.size = Pt(8)
        r.font.color.rgb = _NOTE_COLOR


def _lines(doc, block: Block) -> None:
    for line in block.lines:
        p = _para(doc, after=3)
        run = p.add_run(line)
        run.font.name = _MONO
        run.font.size = Pt(9)
    # One blank line after every format, separating it from the next.
    _para(doc, after=6)


def _typed_section(doc, title: str, type_label: str, blocks: list[Block]) -> None:
    """A '# CDRs'/'# IMEIs' section: one '## <Network> <TYPE>' block per network."""
    if not blocks:
        return
    _heading(doc, title, size=16, color=_SECTION_COLOR, before=22, after=2)
    for network in SECTION_NETWORKS:
        net_blocks = [b for b in blocks if b.network == network]
        if not net_blocks:
            continue
        _heading(
            doc,
            f"{_NETWORK_NAMES[network]} {type_label}",
            size=12,
            color=_NETWORK_COLORS.get(network, _DEFAULT_HEADING_COLOR),
            before=14,
            after=2,
        )
        for block in net_blocks:
            _subtitle(doc, block)
            _lines(doc, block)

    # CDR/IMEI numbers whose network is blank or unrecognised: keep them in this
    # section (listed, not formatted) so no IMEI/CDR is lost, but flag the gap.
    orphan = [b for b in blocks if b.network not in SECTION_NETWORKS]
    if orphan:
        _heading(doc, f"{type_label} — network not set", size=12,
                 color=_DEFAULT_HEADING_COLOR, before=14, after=2)
        for block in orphan:
            _subtitle(doc, block)
            _lines(doc, block)


def _flat_section(doc, title: str, blocks: list[Block]) -> None:
    """Gateway / leftover rows: no per-network grouping, just each block's sets."""
    if not blocks:
        return
    _heading(doc, title, size=16, color=_SECTION_COLOR, before=22, after=2)
    for block in blocks:
        # Leftover ("Other") blocks keep their own descriptive heading; Gateway
        # blocks share the section title, so don't repeat it.
        if block.heading and block.heading != title and block.heading != "Gateway":
            _heading(
                doc, block.heading, size=12,
                color=_NETWORK_COLORS.get(block.network, _DEFAULT_HEADING_COLOR),
                before=14, after=2,
            )
        _subtitle(doc, block)
        _lines(doc, block)


def build_export_document(rows: list[tuple], today: date | None = None):
    """rows are (network, request_type, days, number[, window]) - see formats.build_document."""
    stamp = today or date.today()
    sections = build_document(rows, today=stamp)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    _heading(doc, f"Request Export — {stamp.strftime('%d/%m/%Y')}",
             size=18, color=_SECTION_COLOR, before=0, after=4)

    _typed_section(doc, "CDRs", "CDR", sections.get("CDR") or [])
    _typed_section(doc, "IMEIs", "IMEI", sections.get("IMEI") or [])
    _flat_section(doc, "Gateway", sections.get("GATEWAY") or [])
    _flat_section(doc, "Other (handle manually)", sections.get("MANUAL") or [])

    return doc
